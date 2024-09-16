import sqlite3
from docx import Document

test_doc = Document('test.docx')
con = sqlite3.connect('db.sqlite')
cur = con.cursor()

for index in range(len(test_doc.paragraphs)):
    if index % 2 == 0:
        cur.execute(f'INSERT INTO questions (question, right_answer) VALUES ({test_doc.paragraphs[index].text}, "1")')
        con.commit()
    else:
        cur.execute(f'INSERT INTO questions (question, right_answer) VALUES ({test_doc.paragraphs[index].text}, "2")')
        con.commit()

con.close()